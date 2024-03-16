from docx import Document
from datetime import datetime
import inflect
import comtypes.client
import os

def convert_docx_to_pdf(input_file, output_file):
    doc=Document(input_file)
    word=comtypes.client.CreateObject("Word.Application")
    docx_path=os.path.abspath(input_file)
    pdf_path=os.path.abspath(output_file)

    pdf_format=17
    word.Visible=False
    in_file=word.Documents.Open(docx_path)
    in_file.SaveAs(pdf_path, FileFormat=pdf_format)
    in_file.Close()
    word.Quit()

def fill_invoice_fields(json_data, template_path, output_path):
    p = inflect.engine()
    json_data["invoice_date"] = datetime.now().strftime("%d/%m/%Y")

    total_amount = json_data["total"]
    json_data["pgc"] = str(round(int(total_amount) * 0.0236, 2))
    json_data["subtotal"] = str(round((int(total_amount) - float(json_data["pgc"])) / 1.18, 2))
    json_data["cgst"] = str(round(float(json_data["subtotal"]) * 0.09, 2))
    json_data["sgst"] = json_data["cgst"]
    json_data["amount"] = str(float(json_data["subtotal"]) + int(json_data["discount"]))
    json_data["rate"] = json_data["amount"]
    json_data["total_in_words"] = (p.number_to_words(int(total_amount)) + " only").upper()
    if float(json_data["cgst"]) > 0 and float(json_data["sgst"]) > 0:
        json_data["igst"] = str(0)

    # Generating Invoice Number
    generated_date = json_data["invoice_date"]
    given_date = datetime.strptime(generated_date, "%d/%m/%Y")
    current_year = datetime.now().year
    start_date = datetime(current_year, 4, 1)
    end_date = datetime(current_year + 1, 3, 31)
    id = json_data["id"]
    if start_date <= given_date <= end_date:
        json_data["invoice_no"] = f"{current_year}-{(current_year + 1) % 100}/{given_date.month}/D/00{id}"
    else:
        json_data["invoice_no"] = f"{current_year - 1}-{(current_year) % 100}/{given_date.month}/D/00{id}"
    print(json_data)

    # Load the Word document template
    doc = Document(template_path)

    # Define the mapping between JSON keys and document fields
    # Please dont change the keys of following fields dictionary, as the keys points to the specific cell of table. If you change the keys in the dictionary then make sure to change the same key in the template as well or vice versa
    fields = {
        "Invoice No": "invoice_no",
        "Invoice Date": "invoice_date",
        "Name": "name",
        "Address": "address",
        "GSTIN": "gstin",
        "Product Description": "product_description",
        "Rate": "rate",
        "Amount": "amount",
        "Discount": "discount",
        "Subtotal": "subtotal",
        "CGST (in%)": "cgst",
        "SGST (in%)": "sgst",
        "IGST (in%)": "igst",
        "Payment gateway charges": "pgc",
        "Total (in words)": "total_in_words",
        "Total": "total"
    }

    # Fill in the document fields with JSON values
    for i, table in enumerate(doc.tables):
        for k, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text in fields and (i == 0 or i == 1):
                    row.cells[j + 1].text = json_data[fields[cell.text]]

                if cell.text in fields and i == 2:
                    next_row = table.rows[k + 1]
                    required_cell = next_row.cells[j]
                    required_cell.text = json_data[fields[cell.text]]

                if cell.text in fields and i == 3:
                    if k == 7 and cell.text == "Total (in words)":
                        next_row = table.rows[k + 1]
                        required_cell = next_row.cells[j]
                        required_cell.text = json_data[fields[cell.text]]
                    elif k == 7 and cell.text == "Total":
                        row.cells[j + 1].text = json_data[fields[cell.text]]
                    else:
                        row.cells[j + 2].text = json_data[fields[cell.text]]

    # Save the filled-in document
 

    doc.save(output_path)
    convert_docx_to_pdf(output_path, "output_pdf.pdf")
    print(f"Filled-in document saved to {output_path}")

# Example JSON data
json_data = {
    "name": "Prathamesh Satardekar",
    "address": "110, Shivshakti Industrial Estate, LBS Marg, Opp. Shantiniketan Hospital, Ghatkopar (W), Mumbai - 400086",
    "gstin": "27AACCV9053F1Z6",
    "product_description": "DPL REGISTRATION FEE\nName: DPL - Saudi Arabia\nDate: Dec 2023\nVenue: Saudi Arabia\nRegistration No.: DPL23PG0100001",
    "discount": "0",
    "total": "700",
    "id": "4"
}

# Path to the Word document template and output document
template_path = "input_template.docx"  # Adjust this path to your template file
output_path = "output_template.docx"  # Adjust this path for output file

# Fill in the invoice fields with JSON values
fill_invoice_fields(json_data, template_path, output_path)
